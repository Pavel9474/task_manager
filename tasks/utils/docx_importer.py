import re
from docx import Document
from datetime import datetime
from ..models import Subtask, Employee, ResearchProduct

class ResearchDocxImporter:
    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = Document(file_path)
        print(f"Загружен документ: {file_path}")
        print(f"Количество параграфов: {len(self.doc.paragraphs)}")
        print(f"Количество таблиц: {len(self.doc.tables)}")
    
    def parse_research_task(self):
        """Парсинг ТЗ и возврат структуры данных"""
        result = {
            'title': self._extract_title(),
            'tz_number': self._extract_tz_number(),
            'foundation': self._extract_foundation(),
            'funding_source': self._extract_funding_source(),
            'government_work_name': self._extract_government_work_name(),
            'customer': self._extract_customer(),
            'executor': self._extract_executor(),
            'executor_address': self._extract_executor_address(),
            'start_date': None,
            'end_date': None,
            'location': self._extract_location(),
            'funding_years': self._extract_funding_years(),  # Словарь {год: сумма}
            'goals': self._extract_goals(),
            'tasks': self._extract_tasks(),
            'stages': self._extract_stages()
        }
        
        # Извлекаем даты
        start_date, end_date = self._extract_dates_from_text()
        result['start_date'] = start_date
        result['end_date'] = end_date
        
        print(f"Найдено название: {result['title']}")
        print(f"Найдено шифр: {result['tz_number']}")
        print(f"Найдено финансирование по годам: {result['funding_years']}")
        print(f"Найдено этапов: {len(result['stages'])}")
        
        return result
    def _extract_tz_number(self):
        """Извлечение шифра из текста"""
        for paragraph in self.doc.paragraphs:
            if 'шифр:' in paragraph.text.lower():
                match = re.search(r'шифр:\s*[«"](.+?)[»"]', paragraph.text, re.IGNORECASE)
                if match:
                    return match.group(1)
        return ''

    def _extract_foundation(self):
        """Извлечение основания для проведения работы"""
        for paragraph in self.doc.paragraphs:
            if '1. Основание для проведения работы:' in paragraph.text:
                # Берем следующий абзац
                idx = self.doc.paragraphs.index(paragraph)
                if idx + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[idx + 1].text.strip()
        return ''

    def _extract_funding_source(self):
        """Извлечение источника финансирования"""
        for paragraph in self.doc.paragraphs:
            if '2. Источник финансирования:' in paragraph.text:
                idx = self.doc.paragraphs.index(paragraph)
                if idx + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[idx + 1].text.strip()
        return ''

    def _extract_government_work_name(self):
        """Извлечение наименования государственной работы"""
        for paragraph in self.doc.paragraphs:
            if '3. Наименование государственной работы:' in paragraph.text:
                idx = self.doc.paragraphs.index(paragraph)
                if idx + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[idx + 1].text.strip()
        return ''

    def _extract_customer(self):
        """Извлечение государственного заказчика"""
        for paragraph in self.doc.paragraphs:
            if '4. Государственный заказчик:' in paragraph.text:
                idx = self.doc.paragraphs.index(paragraph)
                if idx + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[idx + 1].text.strip()
        return ''

    def _extract_executor(self):
        """Извлечение исполнителя"""
        for paragraph in self.doc.paragraphs:
            if '5. Исполнитель' in paragraph.text:
                # Берем следующую строку
                idx = self.doc.paragraphs.index(paragraph)
                if idx + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[idx + 1].text.strip()
        return ''

    def _extract_executor_address(self):
        """Извлечение адреса исполнителя"""
        for paragraph in self.doc.paragraphs:
            if '5. Исполнитель' in paragraph.text:
                # Ищем в этом же абзаце после двоеточия
                match = re.search(r'5\. Исполнитель.*?:?\s*(.+?)(?:\n|$)', paragraph.text, re.DOTALL)
                if match:
                    return match.group(1).strip()
        return ''

    def _extract_location(self):
        """Извлечение места проведения работы"""
        for paragraph in self.doc.paragraphs:
            if '7. Место проведения работы:' in paragraph.text:
                idx = self.doc.paragraphs.index(paragraph)
                if idx + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[idx + 1].text.strip()
        return ''

    def _extract_funding_years(self):
        """Извлечение финансирования по годам"""
        funding = {}
        
        # Проходим по всем таблицам
        for table in self.doc.tables:
            for row in table.rows:
                # Получаем текст всех ячеек
                cells_text = [cell.text.strip() for cell in row.cells]
                if not cells_text:
                    continue
                
                # Объединяем текст всех ячеек
                full_text = ' '.join(cells_text)
                
                # Ищем паттерн: "В 2025 году: 44 747 500"
                # Или "8.1 В 2025 году: 44 747 500"
                match = re.search(r'В\s+(\d{4})\s+году:\s+([\d\s]+?)(?:\s*\(|$)', full_text)
                if match:
                    year = int(match.group(1))
                    amount_str = match.group(2).strip()
                    # Убираем пробелы
                    amount_clean = amount_str.replace(' ', '')
                    if amount_clean and amount_clean.isdigit():
                        try:
                            funding[year] = float(amount_clean)
                            print(f"  Найдено: {year} год - {amount_clean} руб.")
                        except:
                            pass
        
        return funding
    def _extract_goals(self):
        """Извлечение целей работы"""
        goals = []
        collecting = False
        for paragraph in self.doc.paragraphs:
            if '9. Цели работы:' in paragraph.text:
                collecting = True
                continue
            if collecting:
                if '10. Задачи работы:' in paragraph.text:
                    break
                text = paragraph.text.strip()
                if text and not text.startswith('----'):
                    goals.append(text)
        return '\n'.join(goals)

    def _extract_tasks(self):
        """Извлечение задач работы"""
        tasks = []
        collecting = False
        for paragraph in self.doc.paragraphs:
            if '10. Задачи работы:' in paragraph.text:
                collecting = True
                continue
            if collecting:
                if '11. Этапы' in paragraph.text:
                    break
                text = paragraph.text.strip()
                if text and not text.startswith('----'):
                    # Убираем номера задач
                    text = re.sub(r'^\d+\.\d+\.?\s*', '', text)
                    tasks.append(text)
        return '\n'.join(tasks)

    def _extract_dates_from_text(self):
        """Извлечение дат начала и окончания из текста"""
        for paragraph in self.doc.paragraphs:
            if '6. Сроки выполнения работы:' in paragraph.text:
                match = re.search(r'начало\s*-\s*(\d{2}\.\d{2}\.\d{4})\s*;\s*окончание\s*-\s*(\d{2}\.\d{2}\.\d{4})', paragraph.text)
                if match:
                    from datetime import datetime
                    start_str, end_str = match.groups()
                    start_date = datetime.strptime(start_str, '%d.%m.%Y').date()
                    end_date = datetime.strptime(end_str, '%d.%m.%Y').date()
                    return start_date, end_date
        return None, None
    def _extract_title(self):
        """Извлечение названия ТЗ"""
        for paragraph in self.doc.paragraphs:
            if 'по теме:' in paragraph.text:
                match = re.search(r'по теме:\s*«(.+?)»', paragraph.text)
                if match:
                    return match.group(1)
                # Альтернативный формат
                match = re.search(r'по теме:\s*"(.+?)"', paragraph.text)
                if match:
                    return match.group(1)
        return None
    
    def _extract_stages(self):
        """Извлечение этапов из таблицы"""
        stages = []
        
        # Ищем таблицу с этапами
        target_table = None
        for table_idx, table in enumerate(self.doc.tables):
            if len(table.rows) > 15 and len(table.columns) == 4:
                target_table = table
                print(f"\n✅ Найдена таблица с этапами (индекс {table_idx + 1})")
                break
        
        if not target_table:
            print("❌ Таблица с этапами не найдена")
            return stages
        
        print(f"Анализ таблицы этапов:")
        print(f"  Всего строк: {len(target_table.rows)}")
        
        current_stage = None
        
        for row_idx, row in enumerate(target_table.rows):
            if row_idx < 2:  # Пропускаем первые две строки (заголовки)
                continue
                
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 4:
                continue
                
            stage_num = cells[0].strip()
            stage_title = cells[1].strip()
            stage_products = cells[2].strip() if len(cells) > 2 else ''
            stage_dates = cells[3].strip() if len(cells) > 3 else ''
            
            # Пропускаем пустые строки
            if not stage_num and not stage_title:
                continue
            
            print(f"\n  Строка {row_idx + 1}: номер='{stage_num}', даты='{stage_dates}'")
            
            # Парсим даты
            start_date, end_date = self._parse_dates(stage_dates)
            
            # Проверяем, является ли строка этапом (целое число)
            if stage_num.isdigit():
                stage_number = int(stage_num)
                
                # Пропускаем строки с номерами колонок
                if stage_title == '2' or stage_title == '3' or stage_title == '4':
                    print(f"    Пропускаем строку-заголовок: {stage_title}")
                    continue
                
                current_stage = {
                    'number': stage_number,
                    'title': stage_title,
                    'products': self._parse_products(stage_products),
                    'start_date': start_date,
                    'end_date': end_date,
                    'substages': []
                }
                stages.append(current_stage)
                print(f"    ✅ Этап {stage_number}: {stage_title[:50]}...")
                if start_date and end_date:
                    print(f"       Даты: {start_date} - {end_date}")
                
            # Проверяем, является ли строка подэтапом (число с точкой)
            elif '.' in stage_num and stage_num.replace('.', '').isdigit():
                if current_stage:
                    stage_main_num = int(stage_num.split('.')[0])
                    if stage_main_num == current_stage['number']:
                        substage = {
                            'number': stage_num,
                            'title': stage_title,
                            'products': self._parse_products(stage_products),
                            'start_date': start_date,
                            'end_date': end_date
                        }
                        current_stage['substages'].append(substage)
                        print(f"    🔹 Подэтап {stage_num} добавлен к этапу {current_stage['number']}")
                        if start_date and end_date:
                            print(f"       Даты: {start_date} - {end_date}")
                    else:
                        print(f"    ⚠️ Подэтап {stage_num} не соответствует текущему этапу {current_stage['number']}, ищем подходящий этап")
                        found = False
                        for stage in stages:
                            if stage['number'] == stage_main_num:
                                substage = {
                                    'number': stage_num,
                                    'title': stage_title,
                                    'products': self._parse_products(stage_products),
                                    'start_date': start_date,
                                    'end_date': end_date
                                }
                                stage['substages'].append(substage)
                                print(f"    🔹 Подэтап {stage_num} добавлен к этапу {stage['number']}")
                                if start_date and end_date:
                                    print(f"       Даты: {start_date} - {end_date}")
                                found = True
                                break
                        if not found:
                            print(f"    ❌ Не найден этап {stage_main_num} для подэтапа {stage_num}")
        
        # Выводим итоговую структуру
        print("\n=== ИТОГОВАЯ СТРУКТУРА ===")
        for stage in stages:
            print(f"Этап {stage['number']}: {stage['title'][:50]}...")
            if stage.get('start_date') and stage.get('end_date'):
                print(f"  Даты: {stage['start_date']} - {stage['end_date']}")
            print(f"  Подэтапов: {len(stage['substages'])}")
            for substage in stage['substages']:
                print(f"    - {substage['number']}: {substage['title'][:50]}...")
                if substage.get('start_date') and substage.get('end_date'):
                    print(f"      Даты: {substage['start_date']} - {substage['end_date']}")
                print(f"      Продукции: {len(substage.get('products', []))}")
        
        return stages

    def _parse_dates(self, date_str):
        """Парсинг строки с датами вида '01.04.2025 - 30.06.2025'"""
        if not date_str or date_str == '-':
            return None, None
        
        try:
            # Ищем паттерн с датами
            import re
            from django.utils import timezone
            from datetime import datetime
            
            date_pattern = r'(\d{2}\.\d{2}\.\d{4})\s*[-—]\s*(\d{2}\.\d{2}\.\d{4})'
            match = re.search(date_pattern, date_str)
            
            if match:
                start_str, end_str = match.groups()
                # Создаем datetime с временем 00:00:00 и делаем его "aware"
                start_date = datetime.strptime(start_str, '%d.%m.%Y').replace(hour=0, minute=0, second=0)
                end_date = datetime.strptime(end_str, '%d.%m.%Y').replace(hour=23, minute=59, second=59)
                
                # Добавляем временную зону
                start_date = timezone.make_aware(start_date)
                end_date = timezone.make_aware(end_date)
                
                return start_date, end_date
        except Exception as e:
            print(f"    Ошибка парсинга дат '{date_str}': {e}")
        
        return None, None
    
    def _parse_products(self, products_str):
        """Парсинг научно-технической продукции"""
        products = []
        if products_str and products_str != '-' and products_str.strip():
            # Разделяем по точке с запятой, точке или новой строке
            parts = re.split(r'[.;\n]\s*', products_str)
            for part in parts:
                part = part.strip()
                # Фильтруем пустые и слишком короткие строки
                if part and len(part) > 5 and not part.startswith(('http', 'www')):
                    # Убираем нумерацию в начале
                    part = re.sub(r'^\d+\.\s*', '', part)
                    products.append(part)
        return products
    
    def create_task_structure(self, task, stages_data):
        """Создание структуры этапов и подэтапов для задачи"""
        from ..models import Subtask, ResearchProduct
        
        print(f"\nСоздание структуры для задачи: {task.title}")
        print(f"Получено этапов: {len(stages_data)}")
        
        # Удаляем все существующие подэтапы и продукцию
        subtasks_to_delete = Subtask.objects.filter(task=task)
        for subtask in subtasks_to_delete:
            # Удаляем продукцию этого подэтапа
            ResearchProduct.objects.filter(subtask=subtask).delete()
        subtasks_to_delete.delete()
        print(f"Удалено старых подэтапов и продукции")
        
        created_count = 0
        created_products_count = 0
        
        # Создаем все этапы и подэтапы
        for stage_idx, stage_data in enumerate(stages_data):
            # Создаем основной этап
            stage_number = str(stage_data['number'])
            
            planned_start = stage_data.get('start_date')
            planned_end = stage_data.get('end_date')
            
            stage = Subtask.objects.create(
                task=task,
                stage_number=stage_number,
                title=stage_data['title'][:200],
                description=f"Этап {stage_data['number']} из ТЗ",
                planned_start=planned_start,
                planned_end=planned_end,
                status='pending',
                priority='medium'
            )
            date_info = f", даты: {planned_start} - {planned_end}" if planned_start and planned_end else ""
            print(f"  ✅ Создан этап {stage_number} ID: {stage.id}{date_info}")
            created_count += 1
            
            # Создаем подэтапы
            for substage_data in stage_data.get('substages', []):
                try:
                    substage_number = str(substage_data['number'])
                    
                    substage_start = substage_data.get('start_date')
                    substage_end = substage_data.get('end_date')
                    
                    # Создаем подэтап
                    substage = Subtask.objects.create(
                        task=task,
                        stage_number=substage_number,
                        title=substage_data['title'][:200],
                        description=f"Подэтап {substage_data['number']} из ТЗ",
                        planned_start=substage_start,
                        planned_end=substage_end,
                        status='pending',
                        priority='medium'
                    )
                    date_info = f", даты: {substage_start} - {substage_end}" if substage_start and substage_end else ""
                    print(f"    🔹 Создан подэтап {substage_number} ID: {substage.id}{date_info}")
                    created_count += 1
                    
                    # СОЗДАЁМ ПРОДУКЦИЮ КАК ОТДЕЛЬНЫЕ ОБЪЕКТЫ
                    products_list = substage_data.get('products', [])
                    for product_name in products_list:
                        if product_name and len(product_name) > 5:
                            product = ResearchProduct.objects.create(
                                subtask=substage,
                                name=product_name[:500],  # Ограничиваем длину
                                description=f"Продукция подэтапа {substage_number}",
                                status='pending'
                            )
                            created_products_count += 1
                            print(f"        📄 Создана продукция: {product_name[:80]}...")
                    
                    print(f"        Продукция: {len(products_list)} позиций создано")
                            
                except Exception as e:
                    print(f"    ❌ Ошибка создания подэтапа {substage_data.get('number')}: {e}")
        
        # Проверяем результат
        final_count = Subtask.objects.filter(task=task).count()
        products_count = ResearchProduct.objects.filter(subtask__task=task).count()
        print(f"\n✅ Всего создано подэтапов: {final_count}")
        print(f"✅ Всего создано продукции: {products_count}")
        
        return created_count
    def import_research(self, default_performers=None):
        """Импорт НИР в базу данных (создаёт ResearchTask)"""
        from ..models import ResearchTask, ResearchStage, ResearchSubstage, ResearchProduct, ResearchTaskFunding
        
        data = self.parse_research_task()
        
        print(f"\n=== СОЗДАНИЕ НИР ===")
        print(f"Название: {data['title']}")
        print(f"Шифр: {data.get('tz_number', '')}")
        
        # Создаем НИР
        research_task = ResearchTask.objects.create(
            title=data['title'],
            tz_number=data.get('tz_number', ''),
            foundation=data.get('foundation', ''),
            funding_source=data.get('funding_source', ''),
            government_work_name=data.get('government_work_name', ''),
            customer=data.get('customer', ''),
            executor=data.get('executor', ''),
            executor_address=data.get('executor_address', ''),
            start_date=data.get('start_date'),
            end_date=data.get('end_date'),
            location=data.get('location', ''),
            goals=data.get('goals', ''),
            tasks=data.get('tasks', ''),
        )
        
        print(f"✅ Создана НИР ID: {research_task.id}")
        
        # Сохраняем финансирование
        for year, amount in data.get('funding_years', {}).items():
            ResearchTaskFunding.objects.create(
                research_task=research_task,
                year=year,
                amount=amount
            )
            print(f"  📊 Добавлено финансирование на {year} год: {amount:,.0f} руб.")
        
        # Создаем этапы
        for stage_data in data.get('stages', []):
            stage = ResearchStage.objects.create(
                research_task=research_task,
                stage_number=stage_data['number'],
                title=stage_data['title'],
                start_date=stage_data.get('start_date'),
                end_date=stage_data.get('end_date')
            )
            print(f"  📁 Создан этап {stage.stage_number}: {stage.title[:50]}...")
            
            if default_performers:
                stage.performers.set(default_performers)
            
            # Создаем подэтапы
            for substage_data in stage_data.get('substages', []):
                substage = ResearchSubstage.objects.create(
                    stage=stage,
                    substage_number=substage_data['number'],
                    title=substage_data['title'],
                    description=f"Подэтап {substage_data['number']}",
                    start_date=substage_data.get('start_date'),
                    end_date=substage_data.get('end_date')
                )
                print(f"    🔹 Подэтап {substage.substage_number}: {substage.title[:50]}...")
                
                if default_performers:
                    substage.performers.set(default_performers)
                
                # Создаем продукцию - используем research_substage, а не substage!
                for product_name in substage_data.get('products', []):
                    if product_name and len(product_name) > 5:
                        ResearchProduct.objects.create(
                            research_substage=substage,  # <-- ИСПРАВЛЕНО!
                            name=product_name[:500],
                            description=f"Продукция подэтапа {substage_data['number']}",
                            status='pending'
                        )
                        print(f"        📄 Продукция: {product_name[:60]}...")
        
        print(f"\n✅ Импорт завершён!")
        return research_task